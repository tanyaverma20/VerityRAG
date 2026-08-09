"""
report_generator.py — Structured research report generation.

Pipeline (matches the Q&A pipeline's shape exactly):

    document_ids
      -> per-document hybrid retrieval + rerank (deterministic, no LLM)
      -> global token-budget selection across all documents (no LLM)
      -> ONE report-generation LLM call
      -> Pydantic-validated JSON (schemas.ResearchReport)
      -> deterministic renderers: JSON -> Markdown, JSON -> PDF

Never one call per paper, never a separate comparison call — a single
report call receives ALL selected evidence (grouped by document) at once
and produces the full structured report in one shot.
"""
from __future__ import annotations

import io
import re
from typing import Any

from pydantic import ValidationError
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib import colors

from config import (
    REPORT_CHUNKS_PER_DOC, MAX_REPORT_CONTEXT_TOKENS, MAX_REPORT_ANSWER_TOKENS, GROQ_API_KEY,
)
from retrieval import retrieve, select_within_token_budget, group_by_document
from query_transform import (
    _call_groq_raw, is_rate_limit_error,
    RATE_LIMIT_MESSAGE, GENERATION_FAILED_MESSAGE, NO_DOCUMENTS_MESSAGE,
)
from schemas import ResearchReport, PaperReport, ComparisonReport
from doc_titles import resolve_display_title

# A fixed, broad query designed to surface chunks touching every report
# section (problem, method, data, results, limitations) rather than one
# narrow topic — retrieval/reranking stay entirely non-LLM either way.
REPORT_QUERY = (
    "research problem, motivation, methodology, architecture, datasets, "
    "experiments, key findings, results, limitations, conclusion"
)


REPORT_PROMPT = """You are VerityRAG, a research assistant generating a structured report.

Using ONLY the evidence below — retrieved from the user's currently active
uploaded documents — produce a structured research report. Never use
outside knowledge to fill missing information, and never invent facts,
numbers, methodologies, datasets, results, or limitations that are not in
the evidence.

{comparison_note}

DOCUMENT ID -> FILENAME:
{doc_filename_map}

EVIDENCE (grouped by document):
{evidence_text}

INSTRUCTIONS:
1. For each document listed above, fill in every field using its exact
   document_id from the map: overview (abstract-level summary),
   main_contribution, methodology, architecture, datasets, evaluation_metrics
   (the metric NAMES used, e.g. "BLEU", "F1"), key_results (the actual
   headline outcomes/numbers), important_calculations (specific formulas or
   computed values mentioned, if any), limitations, final_summary.
2. If the evidence does not cover a field for a document, use null (for
   strings) or an empty list [] (for lists) — never invent it, and never
   pad it with "Not available" text yourself; the application fills that in
   for empty/null fields. Set the report's top-level "evidence_sufficient"
   to false if this happens for most fields.
3. {comparison_instruction}
4. Keep prose concise — this is a report, not a full transcript. Do not
   include chain-of-thought, chunk IDs, or internal technical identifiers.

Output ONLY a compact JSON object matching exactly this schema:
{{
  "title": "report title",
  "overview": "1-3 sentence overview of what this report covers",
  "papers": [
    {{
      "document_id": "...",
      "title": "paper title or null",
      "overview": "... or null",
      "main_contribution": "... or null",
      "methodology": "... or null",
      "architecture": "... or null",
      "datasets": ["..."],
      "evaluation_metrics": ["..."],
      "key_results": ["..."],
      "important_calculations": ["..."],
      "limitations": ["..."],
      "final_summary": "... or null"
    }}
  ],
  "comparison": {comparison_schema},
  "conclusion": "1-3 sentence conclusion",
  "evidence_sufficient": true or false
}}

JSON:
"""

_COMPARISON_SCHEMA = (
    '{"commonalities": ["..."], "differences": ["..."], '
    '"strengths": ["..."], "limitations": ["..."]}'
)


def _gather_report_evidence(document_ids: list[str]) -> dict[str, list[dict]]:
    """
    Deterministic, LLM-free evidence gathering. Retrieves broadly per
    document (so every paper gets real representation regardless of how
    many documents are in scope), then applies ONE global token budget
    across the combined set — never "send the whole PDF", never unbounded
    context growth as more papers are added.
    """
    all_chunks: list[dict] = []
    for doc_id in document_ids:
        chunks = retrieve(
            REPORT_QUERY,
            strategy="hybrid",
            document_ids=[doc_id],
            top_k=REPORT_CHUNKS_PER_DOC,
            apply_parent_context=True,
            apply_token_budget=False,  # per-doc budget deferred to the global pass below
        )
        all_chunks.extend(chunks)

    if not all_chunks:
        return {}

    budgeted = select_within_token_budget(
        all_chunks, max_tokens=MAX_REPORT_CONTEXT_TOKENS, max_per_doc=REPORT_CHUNKS_PER_DOC,
    )
    return group_by_document(budgeted)


def _format_evidence_for_prompt(grouped: dict[str, list[dict]]) -> tuple[str, dict[str, str]]:
    parts = []
    filename_map: dict[str, str] = {}
    for doc_id, chunks in grouped.items():
        source = chunks[0].get("metadata", {}).get("source", doc_id) if chunks else doc_id
        filename_map[doc_id] = source
        parts.append(f"\n--- Document ID: {doc_id} (file: {source}) ---")
        for c in chunks:
            parts.append(c.get("parent_context") or c.get("text", ""))
    return "\n".join(parts), filename_map


def generate_report(document_ids: list[str]) -> dict[str, Any]:
    """
    Returns a dict shaped like:
      {"ok": True, "report": ResearchReport, "citations": [...]}
    or
      {"ok": False, "error": "<clean user-facing message>"}

    Exactly ONE LLM call on the success path (see _call_groq_raw ->
    with_model_fallback: at most one additional fallback attempt on a
    genuinely temporary failure — never more).
    """
    if not document_ids:
        return {"ok": False, "error": NO_DOCUMENTS_MESSAGE}

    grouped = _gather_report_evidence(document_ids)
    if not grouped:
        return {"ok": False, "error": "The uploaded documents do not contain enough information to generate a report."}

    evidence_text, filename_map = _format_evidence_for_prompt(grouped)
    doc_filename_map_text = "\n".join(f"- {doc_id}: {name}" for doc_id, name in filename_map.items())

    is_comparison = len(grouped) > 1
    comparison_note = (
        "This is a COMPARATIVE report across multiple papers." if is_comparison
        else "This report covers a single paper."
    )
    comparison_instruction = (
        "Also fill \"comparison\" with commonalities, differences, strengths, and "
        "limitations ACROSS the papers, grounded only in the evidence above."
        if is_comparison else
        "Set \"comparison\" to null — there is only one paper."
    )
    comparison_schema = _COMPARISON_SCHEMA if is_comparison else "null"

    prompt = REPORT_PROMPT.format(
        comparison_note=comparison_note,
        doc_filename_map=doc_filename_map_text,
        evidence_text=evidence_text,
        comparison_instruction=comparison_instruction,
        comparison_schema=comparison_schema,
    )

    if not GROQ_API_KEY:
        return {"ok": False, "error": GENERATION_FAILED_MESSAGE}

    try:
        raw = _call_groq_raw(prompt, max_tokens=MAX_REPORT_ANSWER_TOKENS)  # <-- the ONE LLM call
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if not match:
            raise ValueError("Model response did not contain a JSON object")
        report = ResearchReport.model_validate_json(match.group())
    except ValidationError:
        return {"ok": False, "error": GENERATION_FAILED_MESSAGE}
    except Exception as e:
        return {"ok": False, "error": RATE_LIMIT_MESSAGE if is_rate_limit_error(e) else GENERATION_FAILED_MESSAGE}

    # Display-layer only: normalize every paper's title to something a human
    # can read — filename first, then the model's own title from this SAME
    # call (no extra LLM call), never the raw document_id. document_id on
    # each PaperReport is untouched; this only rewrites the display title.
    for i, p in enumerate(report.papers):
        p.title = resolve_display_title(p.document_id, filename_map.get(p.document_id), fallback_title=p.title, index=i)

    # Citations stay evidence-derived (deterministic), same rule as Q&A.
    citations = [
        {
            "document_id": doc_id,
            "source": filename_map.get(doc_id, doc_id),
            "text": c.get("text", ""),
            "page_number": c.get("metadata", {}).get("page_number", ""),
        }
        for doc_id, chunks in grouped.items()
        for c in chunks
    ]

    return {"ok": True, "report": report, "citations": citations}


# ---------------------------------------------------------------------------
# Renderers — pure Python, deterministic. The LLM only ever produces the
# JSON; formatting into Markdown/PDF/DOCX happens entirely in application
# code (Feature 7 explicitly requires the LLM never generate the file
# formatting itself).
# ---------------------------------------------------------------------------

UNAVAILABLE_TEXT = "Not available in the uploaded evidence"

# (attribute, section label) for every per-paper field, in report order.
_SCALAR_FIELDS = [
    ("overview", "Overview"),
    ("main_contribution", "Main Contribution"),
    ("methodology", "Methodology"),
    ("architecture", "Architecture"),
]
_LIST_FIELDS = [
    ("datasets", "Datasets"),
    ("evaluation_metrics", "Evaluation Metrics"),
    ("key_results", "Key Results"),
    ("important_calculations", "Important Calculations"),
    ("limitations", "Limitations"),
]
_COMPARISON_FIELDS = [
    ("commonalities", "Commonalities"), ("differences", "Differences"),
    ("strengths", "Strengths"), ("limitations", "Limitations"),
]


def render_report_markdown(report: ResearchReport) -> str:
    lines = [f"# {report.title}", "", report.overview, ""]

    for p in report.papers:
        lines.append(f"## {p.title or 'Untitled Document'}")
        for attr, label in _SCALAR_FIELDS:
            lines += [f"**{label}:** {getattr(p, attr) or UNAVAILABLE_TEXT}", ""]
        for attr, label in _LIST_FIELDS:
            values = getattr(p, attr)
            lines.append(f"**{label}:**")
            lines += [f"- {v}" for v in values] if values else [f"- {UNAVAILABLE_TEXT}"]
            lines.append("")
        lines += [f"**Final Summary:** {p.final_summary or UNAVAILABLE_TEXT}", ""]

    if report.comparison:
        lines.append("## Comparison")
        for field_name, label in _COMPARISON_FIELDS:
            values = getattr(report.comparison, field_name)
            if values:
                lines.append(f"**{label}:**")
                lines += [f"- {v}" for v in values]
                lines.append("")

    lines += ["## Conclusion", report.conclusion, ""]

    if not report.evidence_sufficient:
        lines.append(
            "_Note: the uploaded documents did not fully cover this report — "
            "some sections are marked as not available rather than guessed._"
        )

    return "\n".join(lines)


def render_report_pdf(report: ResearchReport) -> bytes:
    """Renders the validated JSON directly to PDF bytes via reportlab —
    no LLM involvement, no re-generation."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("VRH1", parent=styles["Heading1"], spaceAfter=10)
    h2 = ParagraphStyle("VRH2", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6, textColor=colors.HexColor("#4f46e5"))
    body = ParagraphStyle("VRBody", parent=styles["BodyText"], spaceAfter=8, alignment=TA_LEFT)
    label = ParagraphStyle("VRLabel", parent=styles["BodyText"], fontName="Helvetica-Bold", spaceBefore=6, spaceAfter=2)
    note = ParagraphStyle("VRNote", parent=styles["BodyText"], textColor=colors.grey, fontSize=9)

    def esc(text: str) -> str:
        return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = [Paragraph(esc(report.title), h1), Paragraph(esc(report.overview), body), Spacer(1, 6)]

    for p in report.papers:
        story.append(Paragraph(esc(p.title or 'Untitled Document'), h2))
        for attr, field_label in _SCALAR_FIELDS:
            story += [Paragraph(field_label, label), Paragraph(esc(getattr(p, attr) or UNAVAILABLE_TEXT), body)]
        for attr, field_label in _LIST_FIELDS:
            values = getattr(p, attr)
            story.append(Paragraph(field_label, label))
            items = values if values else [UNAVAILABLE_TEXT]
            story.append(ListFlowable([ListItem(Paragraph(esc(v), body)) for v in items], bulletType="bullet"))
        story += [Paragraph("Final Summary", label), Paragraph(esc(p.final_summary or UNAVAILABLE_TEXT), body)]

    if report.comparison:
        story.append(Paragraph("Comparison", h2))
        for field_name, field_label in _COMPARISON_FIELDS:
            values = getattr(report.comparison, field_name)
            if values:
                story.append(Paragraph(field_label, label))
                story.append(ListFlowable([ListItem(Paragraph(esc(v), body)) for v in values], bulletType="bullet"))

    story.append(Paragraph("Conclusion", h2))
    story.append(Paragraph(esc(report.conclusion), body))

    if not report.evidence_sufficient:
        story.append(Spacer(1, 10))
        story.append(Paragraph(
            "Note: the uploaded documents did not fully cover this report — some sections are marked as not available rather than guessed.",
            note,
        ))

    doc.build(story)
    return buf.getvalue()


def render_report_docx(report: ResearchReport) -> bytes:
    """Renders the validated JSON directly to a .docx via python-docx —
    same deterministic, LLM-free rendering as Markdown/PDF."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(report.title, level=0)
    doc.add_paragraph(report.overview)

    accent = RGBColor(0x4F, 0x46, 0xE5)

    for p in report.papers:
        h = doc.add_heading(p.title or "Untitled Document", level=1)
        h.runs[0].font.color.rgb = accent

        for attr, field_label in _SCALAR_FIELDS:
            doc.add_paragraph().add_run(field_label).bold = True
            doc.add_paragraph(getattr(p, attr) or UNAVAILABLE_TEXT)
        for attr, field_label in _LIST_FIELDS:
            doc.add_paragraph().add_run(field_label).bold = True
            values = getattr(p, attr)
            for v in (values if values else [UNAVAILABLE_TEXT]):
                doc.add_paragraph(v, style="List Bullet")
        doc.add_paragraph().add_run("Final Summary").bold = True
        doc.add_paragraph(p.final_summary or UNAVAILABLE_TEXT)

    if report.comparison:
        h = doc.add_heading("Comparison", level=1)
        h.runs[0].font.color.rgb = accent
        for field_name, field_label in _COMPARISON_FIELDS:
            values = getattr(report.comparison, field_name)
            if values:
                doc.add_paragraph().add_run(field_label).bold = True
                for v in values:
                    doc.add_paragraph(v, style="List Bullet")

    h = doc.add_heading("Conclusion", level=1)
    h.runs[0].font.color.rgb = accent
    doc.add_paragraph(report.conclusion)

    if not report.evidence_sufficient:
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(
            "Note: the uploaded documents did not fully cover this report — "
            "some sections are marked as not available rather than guessed."
        )
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
